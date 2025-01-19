import streamlit as st
import pandas as pd
from docx import Document
from docxtpl import DocxTemplate
# from docx.shared import Inches
# import matplotlib.pyplot as plt
import io
from PIL import Image
from docx.shared import Mm




def main():
    st.title("MEMORIAS DE CÁLCULO")
    st.write("Introducir la información requerida para la generación de la memoria.")

    #Introduccion de plantilla
    template_file = st.file_uploader("Cargar plantilla de memoria", type="docx")

    if template_file:
        st.success("Plantilla cargada correctamente")
        #Introducción de datos
        st.markdown("##### Tipo de normas")
        tipo_norma_sismo_col, tipo_norma_viento_col = st.columns(2)

        tipo_norma_sismo=tipo_norma_sismo_col.selectbox("Tipo de norma (Sismo)", ["NTC","CFE", "Otra"])
        tipo_norma_viento=tipo_norma_viento_col.selectbox("Tipo de norma (Viento)", ["CFE", "Otra"])
        
        # Inicializamos el estado en la sesión
        if "step" not in st.session_state:
            st.session_state.step = 1  # Empieza en el paso 1

        # Función para manejar la navegación entre pasos
        def go_to_step(step):
            st.session_state.step = step

            st.write(f"estamos en el paso {step}")
        
        # Paso 1: Datos Generales
        if st.session_state.step == 1:
            st.header("Paso 1: Datos Generales")
            proyecto = st.text_input("Nombre del proyecto")
            fecha = st.date_input("Fecha")
            autor_col, inicialesautor_col = st.columns([0.7, 0.3])
            autor = autor_col.text_input("Autor")
            inicialesautor = inicialesautor_col.text_input("Iniciales del autor")
            revisor_col, inicialesrevisor_col = st.columns([0.7, 0.3])
            revisor = revisor_col.text_input("Revisor")
            inicialesrevisor = inicialesrevisor_col.text_input("Iniciales del revisor")
            direccion = st.text_input("Dirección del proyecto")
            st.markdown("###### Coordenadas")
            latitud_col,longitud_col=st.columns(2)
            latitud=latitud_col.text_input("Latitud")
            longitud=longitud_col.text_input("Longitud")

            if st.button("Siguiente"):
                st.session_state.proyecto = proyecto
                st.session_state.fecha = fecha
                st.session_state.autor = autor
                st.session_state.inicialesautor = inicialesautor
                st.session_state.revisor = revisor
                st.session_state.inicialesrevisor = inicialesrevisor
                st.session_state.direccion = direccion
                st.session_state.latitud = latitud
                st.session_state.longitud = longitud
                go_to_step(2)

        # Paso 2: Descripciones
        elif st.session_state.step == 2:
            st.header("Paso 2: Descripciones")
            descripcionarq = st.text_area("Descripción arquitectónica")
            st.markdown("###### Descripción de la estructura")
            descripciongrav = st.text_area("Descripción del sistema gravitacional")
            descripsistemlat = st.text_area("Descripción del sistema lateral")
            descripsistemcim = st.text_area("Descripción de la cimentación")

            if st.button("Paso 1"):
                go_to_step(1)
            elif st.button("Paso 3"):
                st.session_state.descripcionarq = descripcionarq
                st.session_state.descripciongrav = descripciongrav
                st.session_state.descripsistemlat = descripsistemlat
                st.session_state.descripsistemcim = descripsistemcim
                go_to_step(3)

        # Paso 3: Mecánica de Suelos
        elif st.session_state.step == 3:
            st.header("Paso 3: Mecánica de Suelos")
            autormecsuelos = st.text_input("Despacho que realizó estudio de mecánica de suelos")
            docmecsuelos = st.text_input("Nombre del informe de mecánica de suelos")

            if st.button("Paso 2"):
                go_to_step(2)
            elif st.button("Paso 4"):
                st.session_state.autormecsuelos = autormecsuelos
                st.session_state.docmecsuelos = docmecsuelos
                go_to_step(4)

        # Paso 4: Información Sísmica
        elif st.session_state.step == 4:
            st.header("Paso 4: Información Sísmica")
            #tipo_norma_sismo = st.selectbox("Tipo de norma (Sismo)", ["NTC", "CFE", "Otra"])

            if tipo_norma_sismo == "CFE":
                normasismo = st.selectbox("Norma aplicable para análisis sísmico", ["CFE-MDOCS-2015", "CFE-MDOCS-2008"])
                regionsismoselect = st.selectbox("Región sísmica", ["A", "B", "C", "D"])
                regionsismo = f"la región {regionsismoselect}"
                nivelzonificacion = "la República Mexicana"
                gruposismoselect=st.selectbox("Grupo (Sismo)",["A+","A","B"])
                gruposismo=f"Grupo {gruposismoselect}"
                if gruposismoselect=="A+":
                    descripgruposismo="Estructuras en las que se requiere un grado de seguridad extrema, ya que su falla causaría cientos o miles de víctimas, y/o graves pérdidas y daños económicos, culturales, ecológicos o sociales"
                elif gruposismoselect=="A":
                    descripgruposismo="Estructuras en que se requiere un grado de seguridad alto. Construcciones cuya falla estructural causaría la pérdida de un número elevado de vidas o pérdidas económicas, daños ecológicos o culturales, científicos o tecnológicos de magnitud intensa o excepcionalmente alta, o que constituyan un peligro significativo por contener sustancias tóxicas o inflamables, así como construcciones cuyo funcionamiento sea esencial después de un sismo"
                else:
                    descripgruposismo="Estructuras en las que se requiere un grado de seguridad convencional. Construcciones cuya falla estructural ocasionaría la pérdida de un número reducido de vidas, pérdidas económicas moderadas o pondría en peligro otras construcciones de este grupo y/o daños a las del Grupo A+ y A moderados"
                programasismo="PRODISIS"
            elif tipo_norma_sismo == "Otra":
                normasismo = st.text_input("Indicar norma")
                regionsismoselect = st.text_input("Región sísmica")
                regionsismo = f"la región {regionsismoselect}"
                nivelzonificacion = st.text_input("Entidad (ciudad y/o país)")
                gruposismo=st.text_input("Indicar Grupo")
            elif tipo_norma_sismo == "NTC":
                normasismo = st.selectbox("Norma aplicable para análisis sísmico", ["NTC-Sismo-2004", "NTC-Sismo-2017", "NTC-Sismo-2020", "NTC-Sismo-2023"])
                regionsismoselect = st.selectbox("Zona sísmica", ["I", "II", "III"])
                regionsismo = f"la zona {regionsismoselect}"
                nivelzonificacion = "la Ciudad de México"
                gruposismoselect=st.selectbox("Grupo (Sismo)",["A","B"])
                gruposismo=f"Grupo {gruposismoselect}"
                if gruposismoselect=="A":
                    descripgruposismo="Edificaciones cuya falla estructural podría tener consecuencias particularmente graves"
                else:
                    descripgruposismo="Edificaciones comunes destinadas a viviendas, oficinas y locales comerciales, hoteles y construcciones comerciales e industriales no incluidas en el Grupo A"
                programasismo="SASID"
            else:
                normasismo, regionsismo, nivelzonificacion = None, None, None

            metodosismo=st.text_area("Descripción del método sismico aplicado")
            ductilidad=st.selectbox("Ductilidad",["Q=1", "Q=1.25", "Q=2","Q=3","Q=4"])
            driftpclim=st.text_input("Distorsión límite para prevención de colapso")
            driftservlim=st.text_input("Distorsión límite para servicio")

            if st.button("Paso 3"):
                go_to_step(3)
            elif st.button("Paso 5"):
                st.session_state.normasismo = normasismo
                st.session_state.regionsismo = regionsismo
                st.session_state.nivelzonificacion = nivelzonificacion
                st.session_state.gruposismo = gruposismo
                st.session_state.descripgruposismo = descripgruposismo
                st.session_state.programasismo = programasismo
                st.session_state.metodosismo = metodosismo
                st.session_state.ductilidad = ductilidad
                st.session_state.driftpclim = driftpclim
                st.session_state.driftservlim = driftservlim
                go_to_step(5)

        # Paso 5: Informacion de viento
        elif st.session_state.step == 5:
            st.header("Paso 5: Información de viento")
            if tipo_norma_viento=="CFE":
                normaviento=st.selectbox("Norma aplicable para análisis por viento", ["CFE-MDOCV-2020", "CFE-MDOCV-2008"])
                grupovientoselect=st.selectbox("Grupo (Viento)",["A","B","C"])
                grupoviento=f"Grupo {grupovientoselect}"
                tipoviento=st.selectbox("Clasificación de la estructura ante la acción de viento", ["Tipo 1", "Tipo 2", "Tipo 3", "Tipo 4"])
                categoriarug=st.selectbox("Categoría de rugosidad", ["Categoría 1", "Categoría 2", "Categoría 3", "Categoría 4"])
                if categoriarug=="Categoría 1":
                    alturagrad="245"
                    exponenteviento="0.099"
                elif categoriarug=="Categoría 2":
                    alturagrad="315"
                    exponenteviento="0.128"
                elif categoriarug=="Categoría 3":
                    alturagrad="390"
                    exponenteviento="0.156"
                else:
                    alturagrad="455"
                    exponenteviento="0.170"

                facttopoviento=st.text_input("Factor de topografía")
                tipotopoviento=st.selectbox("Tipo de terreno", ["Valles cerrados", "Terreno plano", "Promontorios", "Terraplenes"])
            else:
                normaviento=st.text_input("Norma aplicable para análisis por viento")
                grupoviento=st.text_input("Indicar grupo")
                tipoviento=st.text_input("Clasificación de la estructura ante la acción de viento")
                categoriarug=st.text_input("Categoría de rugosidad")
                alturagrad=None
                exponenteviento=None
                facttopoviento=st.text_input("Factor de topografía")
                tipotopoviento=st.text_input("Tipo de terreno")

            if st.button("Paso 4"):
                go_to_step(4)
            elif st.button("Paso 6"):
                st.session_state.normaviento = normaviento
                st.session_state.grupoviento = grupoviento
                st.session_state.tipoviento = tipoviento
                st.session_state.categoriarug = categoriarug
                st.session_state.alturagrad = alturagrad
                st.session_state.exponenteviento = exponenteviento
                st.session_state.facttopoviento = facttopoviento
                st.session_state.tipotopoviento = tipotopoviento
                go_to_step(6)

        # Paso 6: Software Utilizado
        elif st.session_state.step == 6:
            st.header("Paso 6: Software Utilizado")
            softwares=st.multiselect("Selecciona los software utilizados", ["ETABS", "SAP2000", "ROBOT", "SAFE", "EXCEL", "RAM Connection", "IDEA Statica Connection"])
            software1=softwares[0] if len(softwares)>0 else None
            software2=softwares[1] if len(softwares)>1 else None
            software3=softwares[2] if len(softwares)>2 else None
            software4=softwares[3] if len(softwares)>3 else None
            software5=softwares[4] if len(softwares)>4 else None
            software6=softwares[5] if len(softwares)>5 else None
            software7=softwares[6] if len(softwares)>6 else None
            
            ###1
            if software1=="SAFE":
                descripsoft1="Análisis y diseño de sistema de piso"
            elif software1=="EXCEL":
                descripsoft1="Diseño de elementos estructurales, conexiones, cálculo de acciones, etc."
            elif software1=="RAM Connection" or software1=="IDEA Statica Connection":
                descripsoft1="Diseño de conexiones"
            elif software1=="ETABS" or software1=="SAP2000" or software1=="ROBOT":
                descripsoft1="Análisis y diseño estructural"
            else:
                descripsoft1=None
            ###2
            if software2=="SAFE":
                descripsoft2="Análisis y diseño de sistema de piso"
            elif software2=="EXCEL":
                descripsoft2="Diseño de elementos estructurales, conexiones, cálculo de acciones, etc."
            elif software2=="RAM Connection" or software2=="IDEA Statica Connection":
                descripsoft2="Diseño de conexiones"
            elif software2=="ETABS" or software2=="SAP2000" or software2=="ROBOT":
                descripsoft2="Análisis y diseño estructural"
            else:
                descripsoft2=None
            ###3
            if software3=="SAFE":
                descripsoft3="Análisis y diseño de sistema de piso"
            elif software3=="EXCEL":
                descripsoft3="Diseño de elementos estructurales, conexiones, cálculo de acciones, etc."
            elif software3=="RAM Connection" or software3=="IDEA Statica Connection":
                descripsoft3="Diseño de conexiones"
            elif software3=="ETABS" or software3=="SAP2000" or software3=="ROBOT":
                descripsoft3="Análisis y diseño estructural"
            else:
                descripsoft3=None
            ###4
            if software4=="SAFE":
                descripsoft4="Análisis y diseño de sistema de piso"
            elif software4=="EXCEL":
                descripsoft4="Diseño de elementos estructurales, conexiones, cálculo de acciones, etc."
            elif software4=="RAM Connection" or software4=="IDEA Statica Connection":
                descripsoft4="Diseño de conexiones"
            elif software4=="ETABS" or software4=="SAP2000" or software4=="ROBOT":
                descripsoft4="Análisis y diseño estructural"
            else:
                descripsoft4=None
            ###5
            if software5=="SAFE":
                descripsoft5="Análisis y diseño de sistema de piso"
            elif software5=="EXCEL":
                descripsoft5="Diseño de elementos estructurales, conexiones, cálculo de acciones, etc."
            elif software5=="RAM Connection" or software5=="IDEA Statica Connection":
                descripsoft5="Diseño de conexiones"
            elif software5=="ETABS" or software5=="SAP2000" or software5=="ROBOT":
                descripsoft5="Análisis y diseño estructural"
            else:
                descripsoft5=None
            ###6
            if software6=="SAFE":
                descripsoft6="Análisis y diseño de sistema de piso"
            elif software6=="EXCEL":
                descripsoft6="Diseño de elementos estructurales, conexiones, cálculo de acciones, etc."
            elif software6=="RAM Connection" or software6=="IDEA Statica Connection":
                descripsoft6="Diseño de conexiones"
            elif software6=="ETABS" or software6=="SAP2000" or software6=="ROBOT":
                descripsoft6="Análisis y diseño estructural"
            else:
                descripsoft6=None
            ###7
            if software7=="SAFE":
                descripsoft7="Análisis y diseño de sistema de piso"
            elif software7=="EXCEL":
                descripsoft7="Diseño de elementos estructurales, conexiones, cálculo de acciones, etc."
            elif software7=="RAM Connection" or software7=="IDEA Statica Connection":
                descripsoft7="Diseño de conexiones"
            elif software7=="ETABS" or software7=="SAP2000" or software7=="ROBOT":
                descripsoft7="Análisis y diseño estructural"
            else:
                descripsoft7=None
            

            if st.button("Paso 5"):
                go_to_step(5)
            elif st.button("Paso 7"):
                st.session_state.software1 = software1
                st.session_state.software2 = software2
                st.session_state.software3 = software3
                st.session_state.software4 = software4
                st.session_state.software5 = software5
                st.session_state.software6 = software6
                st.session_state.software7 = software7

                st.session_state.descripsoft1 = descripsoft1
                st.session_state.descripsoft2 = descripsoft2
                st.session_state.descripsoft3 = descripsoft3
                st.session_state.descripsoft4 = descripsoft4
                st.session_state.descripsoft5 = descripsoft5
                st.session_state.descripsoft6 = descripsoft6
                st.session_state.descripsoft7 = descripsoft7
                go_to_step(7)

        # Paso 7: Configuracion de modelado
        elif st.session_state.step == 7:
            st.header("Paso 7: Consideraciones de Modelado")
            tipolosa=st.selectbox("Sistema de piso modelado como:", ["membrana", "shell"])
            diafragma=st.selectbox("Tipo de driafragma", ["rígido","semirrígido"])
            
            crackcolumn=st.slider("Factor de agrietamiento en columnas",min_value=0.0,max_value=1.0,value=0.7,step=0.05)
            crackbeam=st.slider("Factor de agrietamiento en trabes",min_value=0.0,max_value=1.0,value=0.5,step=0.05)
            crackwall1=st.slider("Factor de agrietamiento en muros (en el plano)",min_value=0.0,max_value=1.0,value=0.5,step=0.05)
            crackwall2=st.slider("Factor de agrietamiento en muros (fuera del plano)",min_value=0.0,max_value=1.0,value=0.25,step=0.05)
            crackcoupbeam=st.slider("Factor de agrietamiento en trabes de acople",min_value=0.0,max_value=1.0,value=0.3,step=0.05)
            crackslab=st.slider("Factor de agrietamiento en losas",min_value=0.0,max_value=1.0,value=0.25,step=0.05)
                
            rigidpanel=st.slider("Rigidez efectiva del panel (%)",min_value=0,max_value=100,value=50,step=10)

            amortiguamiento=st.slider("Amortiguamiento (%)",min_value=0,max_value=50,value=5,step=1)

            if st.button("Paso 6"):
                go_to_step(6)
            elif st.button("Paso 8"):
                st.session_state.tipolosa = tipolosa
                st.session_state.diafragma = diafragma
                st.session_state.crackcolumn = crackcolumn
                st.session_state.crackbeam = crackbeam
                st.session_state.crackwall1 = crackwall1
                st.session_state.crackwall2 = crackwall2
                st.session_state.crackcoupbeam = crackcoupbeam
                st.session_state.crackslab = crackslab
                st.session_state.rigidpanel = rigidpanel
                st.session_state.amortiguamiento = amortiguamiento
                go_to_step(8)

        # Paso 8: Cortante basal
        elif st.session_state.step == 8:
            st.header("Paso 8: Cortante Basal")
            askrevcortantebasal=st.selectbox("¿Fue necesario incrementar fuerzas sísmicas para cumplir con cortante basal mínimo?", ["No", "Si"])
            if askrevcortantebasal=="Si":
                revcortantebasal="De acuerdo con el cálculo realizado, es necesario afectar las ordenadas espectrales para alcanzar el cortante basal mínimo. Cabe aclarar que dichos factores fueron tomados en cuenta para reportar los desplazamientos mostrados en la sección anterior"
            else:
                revcortantebasal=None

            if st.button("Paso 7"):
                go_to_step(7)
            elif st.button("Paso 9"):
                st.session_state.revcortantebasal = revcortantebasal
                go_to_step(9)

        # Paso 9: Resumen
        elif st.session_state.step == 9:
            st.header("Paso 9: Resumen")

            #st.write("### Datos generales")
            proyecto=st.session_state.get('proyecto')
            fecha=st.session_state.get('fecha')
            autor=st.session_state.get('autor')
            inicialesautor=st.session_state.get('inicialesautor')
            revisor=st.session_state.get('revisor')
            inicialesrevisor=st.session_state.get('inicialesrevisor')
            direccion=st.session_state.get('direccion')
            latitud=st.session_state.get('latitud')
            longitud=st.session_state.get('longitud')

            #st.write("### Descripciones")
            descripcionarq=st.session_state.get('descripcionarq')
            descripciongrav=st.session_state.get('descripciongrav')
            descripsistemlat=st.session_state.get('descripsistemlat')
            descripsistemcim=st.session_state.get('descripsistemcim')
            
            #st.write("### Mecánica de suelos")
            autormecsuelos=st.session_state.get('autormecsuelos')
            docmecsuelos=st.session_state.get('docmecsuelos')

            #st.write("### Información sismica")
            normasismo=st.session_state.get('normasismo')
            regionsismo=st.session_state.get('regionsismo')
            nivelzonificacion=st.session_state.get('nivelzonificacion')
            gruposismo=st.session_state.get('gruposismo')
            descripgruposismo=st.session_state.get('descripgruposismo')
            programasismo=st.session_state.get('programasismo')
            metodosismo=st.session_state.get('metodosismo')
            ductilidad=st.session_state.get('ductilidad')
            driftpclim=st.session_state.get('driftpclim')
            driftservlim=st.session_state.get('driftservlim')

            #st.write("### Información viento")
            normaviento=st.session_state.get('normaviento')
            grupoviento=st.session_state.get('grupoviento')
            tipoviento=st.session_state.get('tipoviento')
            categoriarug=st.session_state.get('categoriarug')
            alturagrad=st.session_state.get('alturagrad')
            exponenteviento=st.session_state.get('exponenteviento')
            facttopoviento=st.session_state.get('facttopoviento')
            tipotopoviento=st.session_state.get('tipotopoviento')

            #st.write("### Software")
            software1=st.session_state.get('software1')
            software2=st.session_state.get('software2')
            software3=st.session_state.get('software3')
            software4=st.session_state.get('software4')
            software5=st.session_state.get('software5')
            software6=st.session_state.get('software6')
            software7=st.session_state.get('software7')

            descripsoft1=st.session_state.get('descripsoft1')
            descripsoft2=st.session_state.get('descripsoft2')
            descripsoft3=st.session_state.get('descripsoft3')
            descripsoft4=st.session_state.get('descripsoft4')
            descripsoft5=st.session_state.get('descripsoft5')
            descripsoft6=st.session_state.get('descripsoft6')
            descripsoft7=st.session_state.get('descripsoft7')

            #st.write("### Modelado")
            tipolosa=st.session_state.get('tipolosa')
            diafragma=st.session_state.get('diafragma')
            crackcolumn=st.session_state.get('crackcolumn')
            crackbeam=st.session_state.get('crackbeam')
            crackwall1=st.session_state.get('crackwall1')
            crackwall2=st.session_state.get('crackwall2')
            crackcoupbeam=st.session_state.get('crackcoupbeam')
            crackslab=st.session_state.get('crackslab')
            rigidpanel=st.session_state.get('rigidpanel')
            amortiguamiento=st.session_state.get('amortiguamiento')

            #st.write("### Cortante basal")
            revcortantebasal=st.session_state.get('revcortantebasal')


            ########################
            texto_reemplazado = {'proyecto':proyecto, 'fecha':fecha, 'autor':autor, 'inicialesautor':inicialesautor, 'revisor':revisor, 'inicialesrevisor':inicialesrevisor, 
            'direccion':direccion, 'latitud':latitud, 'longitud':longitud, 'descripcionarq':descripcionarq, 'descripciongrav':descripciongrav, 'descripsistemlat':descripsistemlat,
            'descripsistemcim':descripsistemcim, 'autormecsuelos':autormecsuelos, 'docmecsuelos':docmecsuelos, 'normasismo':normasismo, 'regionsismo':regionsismo,
            'nivelzonificacion':nivelzonificacion, 'gruposismo':gruposismo, 'descripgruposismo':descripgruposismo, 'programasismo':programasismo, 'metodosismo':metodosismo,
            'ductilidad':ductilidad, 'driftpclim':driftpclim, 'driftservlim':driftservlim, 'normaviento':normaviento, 'grupoviento':grupoviento, 'tipoviento':tipoviento,
            'categoriarug':categoriarug, 'alturagrad':alturagrad, 'exponenteviento':exponenteviento, 'facttopoviento':facttopoviento, 'tipotopoviento':tipotopoviento,
            'software1':software1, 'software2':software2, 'software3':software3, 'software4':software4, 'software5':software5, 'software6':software6, 'software7':software7,
            'descripsoft1':descripsoft1, 'descripsoft2':descripsoft2, 'descripsoft3':descripsoft3, 'descripsoft4':descripsoft4, 'descripsoft5':descripsoft5,
            'descripsoft6':descripsoft6, 'descripsoft7':descripsoft7, 'tipolosa':tipolosa, 'diafragma':diafragma, 'crackcolumn':crackcolumn, 'crackbeam':crackbeam,
            'crackwall1':crackwall1, 'crackwall2':crackwall2, 'crackcoupbeam':crackcoupbeam, 'crackslab':crackslab, 'rigidpanel':rigidpanel, 'amortiguamiento':amortiguamiento,
            'revcortantebasal':revcortantebasal}

            if st.button("Paso 8"):
                go_to_step(8)
            elif st.button("Paso 10"):
                
                go_to_step(10)
            

            st.write(texto_reemplazado)

        # Paso 10: Imagenes
        elif st.session_state.step == 10:
            st.header("Paso 10: Imagenes")     

            #Introduccion de figura 1.1 Proyecto
            st.markdown("#### Imagen del proyecto")
            figura1_1 = st.file_uploader("Cargar figura 1.1", type=["jpg", "jpeg", "png"])
            tit_fig_1_1=st.text_input("Titulo de figura 1.1",value="Vista conceptual del edificio")
            # Mostrar la imagen cargada
            if figura1_1 and tit_fig_1_1:
                prefig1_1 = Image.open(figura1_1)
                st.image(prefig1_1, caption=tit_fig_1_1, use_column_width=True)
                image_path = f"temp_{figura1_1.name}"
                with open(image_path, "wb") as f: 
                    f.write(figura1_1.getbuffer())

                

            #Introduccion de figura 1.2 Ubicacion
            st.markdown("#### Imagen de ubicación")
            proyecto=st.session_state.get('proyecto')
            figura1_2 = st.file_uploader("Cargar figura 1.2", type=["jpg", "jpeg", "png"])
            #tit_fig_1_2=st.text_input("Titulo de figura 1.2",value="Vista conceptual del edificio")
            # Mostrar la imagen cargada
            if figura1_2:
                prefig1_2 = Image.open(figura1_2)
                st.image(prefig1_2, caption=f"Ubicación de {proyecto}", use_column_width=True)

            #Introduccion de figura 1.3 a figura 1.22 Arquitectura
            st.markdown("#### Imagenes de arquitectura (Plantas, elevaciones, etc.)")
            num_fig_arq=st.slider("Num. de figuras para descripción de arquitectura",min_value=0,max_value=20,value=5,step=1)

            if num_fig_arq:
                figuras_arq = []
                tits_fig_arq = []
                for i in range(1, num_fig_arq + 1):
                    figura_arq = st.file_uploader(f"Cargar figura 1.{2+i}", type=["jpg", "jpeg", "png"])
                    tit_fig_arq = st.text_input(f"Titulo de figura 1.{2+i}", key=f"fig_arq_{i}")
                    if figura_arq and tit_fig_arq:
                        figuras_arq.append(figura_arq)
                        tits_fig_arq.append(tit_fig_arq)

                        # Mostrar la imagen cargada
                        image = Image.open(figura_arq)
                        st.image(image, caption=tit_fig_arq, use_column_width=True)

            #Introduccion de figura 5.1 plano de cargas
            st.markdown("#### Imagen de planos de cargas")
            figura5_1 = st.file_uploader("Cargar figura 5.1", type=["jpg", "jpeg", "png"])
            tit_fig_5_1=st.text_input("Titulo de figura 5.1",value="Planos de cargas")
            # Mostrar la imagen cargada
            if figura5_1 and tit_fig_5_1:
                prefig5_1 = Image.open(figura5_1)
                st.image(prefig5_1, caption=tit_fig_5_1, use_column_width=True)

            #Introduccion de figura 5.2 regionalizacion
            st.markdown("#### Imagen de regionalización sísmica")
            nivelzonificacion=st.session_state.get('nivelzonificacion')
            figura5_2 = st.file_uploader("Cargar figura 5.2", type=["jpg", "jpeg", "png"])
            # Mostrar la imagen cargada
            if figura5_2:
                prefig5_2 = Image.open(figura5_2)
                st.image(prefig5_2, caption=f"Regionalización símica de {nivelzonificacion}", use_column_width=True)
            
            #Introduccion de figura 5.3 espectro elastico
            st.markdown("#### Imagen de espectro elástico")
            figura5_3 = st.file_uploader("Cargar figura 5.3", type=["jpg", "jpeg", "png"])
            tit_fig_5_3=st.text_input("Titulo de figura 5.3",value="Espectro elástico")
            # Mostrar la imagen cargada
            if figura5_3:
                prefig5_3 = Image.open(figura5_3)
                st.image(prefig5_3, caption=tit_fig_5_3, use_column_width=True)
            
            #Introduccion de figura 5.4 espectro reducido y de servicio
            st.markdown("#### Imagen de espectro de diseño y de servicio")
            figura5_4 = st.file_uploader("Cargar figura 5.4", type=["jpg", "jpeg", "png"])
            tit_fig_5_4=st.text_input("Titulo de figura 5.4",value="Espectro de diseño reducido y de servicio")
            # Mostrar la imagen cargada
            if figura5_4:
                prefig5_4 = Image.open(figura5_4)
                st.image(prefig5_4, caption=tit_fig_5_4, use_column_width=True)

            #Introduccion de figura 5.5 a figura 5.9 viento
            st.markdown("#### Imagenes de viento (Sugerencia: Mapa de velocidad regional)")
            num_fig_viento=st.slider("Num. de figuras para viento",min_value=0,max_value=5,value=1,step=1)

            if num_fig_viento:
                figuras_viento = []
                tits_fig_viento = []
                for i in range(1, num_fig_viento + 1):
                    figura_viento = st.file_uploader(f"Cargar figura 5.{4+i}", type=["jpg", "jpeg", "png"])
                    tit_fig_viento = st.text_input(f"Titulo de figura 5.{4+i}", key=f"fig_viento_{i}")
                    if figura_viento and tit_fig_viento:
                        figuras_viento.append(figura_viento)
                        tits_fig_viento.append(tit_fig_viento)

                        # Mostrar la imagen cargada
                        image = Image.open(figura_viento)
                        st.image(image, caption=tit_fig_viento, use_column_width=True)

            #Introduccion de figura 10.1 a figura 10.5 modelo estructural
            st.markdown("#### Imagenes de modelo estructural")
            num_fig_modest=st.slider("Num. de figuras para modelo estructural",min_value=0,max_value=5,value=1,step=1)

            if num_fig_modest:
                figuras_modest = []
                tits_fig_modest = []
                for i in range(1, num_fig_modest + 1):
                    figura_modest = st.file_uploader(f"Cargar figura 10.{i}", type=["jpg", "jpeg", "png"])
                    tit_fig_modest = st.text_input(f"Titulo de figura 10.{i}", key=f"fig_modest_{i}")
                    if figura_modest and tit_fig_modest:
                        figuras_modest.append(figura_modest)
                        tits_fig_modest.append(tit_fig_modest)

                        # Mostrar la imagen cargada
                        image = Image.open(figura_modest)
                        st.image(image, caption=tit_fig_modest, use_column_width=True)

            #Introduccion de figura 10.6 Configuracion modos de vibrar
            st.markdown("#### Imagenes de los modos de vibrar")
            modo1_col, modo2_col, modo3_col = st.columns(3)
            # Modo 1
            #modo1_col.markdown("##### Imagen del modo 1")
            figura10_6a = modo1_col.file_uploader("Cargar figura de modo 1", type=["jpg", "jpeg", "png"])
            tit_fig_10_6a=modo1_col.text_input("Periodo de modo 1")
            # Modo 2
            #modo2_col.markdown("##### Imagen del modo 2")
            figura10_6b = modo2_col.file_uploader("Cargar figura de modo 2", type=["jpg", "jpeg", "png"])
            tit_fig_10_6b=modo2_col.text_input("Periodo de modo 2")
            # Modo 3
            #modo3_col.markdown("##### Imagen del modo 3")
            figura10_6c = modo3_col.file_uploader("Cargar figura de modo 3", type=["jpg", "jpeg", "png"])
            tit_fig_10_6c=modo3_col.text_input("Periodo de modo 3")
            # Mostrar la imagen cargada
            if figura10_6a and figura10_6b and figura10_6c and tit_fig_10_6a and tit_fig_10_6b and tit_fig_10_6c:
                prefig10_6a = Image.open(figura10_6a)
                prefig10_6b = Image.open(figura10_6b)
                prefig10_6c = Image.open(figura10_6c)
                modo1_col.image(prefig10_6a, caption=f"Modo 1. T={tit_fig_10_6a} s", use_column_width=True)
                modo2_col.image(prefig10_6b, caption=f"Modo 2. T={tit_fig_10_6b} s", use_column_width=True)
                modo3_col.image(prefig10_6c, caption=f"Modo 3. T={tit_fig_10_6c} s", use_column_width=True)

            #Introduccion de figura 10.9 a figura 10.28 deflexiones verticales
            st.markdown("#### Imagenes de deflexiones verticales")
            num_fig_dz=st.slider("Num. de figuras para mostrar delfexiones verticales",min_value=0,max_value=20,value=5,step=1)

            if num_fig_dz:
                figuras_dz = []
                tits_fig_dz = []
                for i in range(1, num_fig_dz + 1):
                    figura_dz = st.file_uploader(f"Cargar figura 10.{num_fig_modest+3+i}", type=["jpg", "jpeg", "png"])
                    tit_fig_dz = st.text_input(f"Titulo de figura 10.{num_fig_modest+3+i}", key=f"fig_dz_{i}")
                    if figura_dz and tit_fig_dz:
                        figuras_dz.append(figura_dz)
                        tits_fig_dz.append(tit_fig_dz)

                        # Mostrar la imagen cargada
                        image = Image.open(figura_dz)
                        st.image(image, caption=tit_fig_dz, use_column_width=True)

            #st.write("### Datos generales")
            proyecto=st.session_state.get('proyecto')
            fecha=st.session_state.get('fecha')
            ano=fecha.year
            mes=fecha.month
            dia=fecha.day
            autor=st.session_state.get('autor')
            inicialesautor=st.session_state.get('inicialesautor')
            revisor=st.session_state.get('revisor')
            inicialesrevisor=st.session_state.get('inicialesrevisor')
            direccion=st.session_state.get('direccion')
            latitud=st.session_state.get('latitud')
            longitud=st.session_state.get('longitud')

            #st.write("### Descripciones")
            descripcionarq=st.session_state.get('descripcionarq')
            descripciongrav=st.session_state.get('descripciongrav')
            descripsistemlat=st.session_state.get('descripsistemlat')
            descripsistemcim=st.session_state.get('descripsistemcim')
            
            #st.write("### Mecánica de suelos")
            autormecsuelos=st.session_state.get('autormecsuelos')
            docmecsuelos=st.session_state.get('docmecsuelos')

            #st.write("### Información sismica")
            normasismo=st.session_state.get('normasismo')
            regionsismo=st.session_state.get('regionsismo')
            nivelzonificacion=st.session_state.get('nivelzonificacion')
            gruposismo=st.session_state.get('gruposismo')
            descripgruposismo=st.session_state.get('descripgruposismo')
            programasismo=st.session_state.get('programasismo')
            metodosismo=st.session_state.get('metodosismo')
            ductilidad=st.session_state.get('ductilidad')
            driftpclim=st.session_state.get('driftpclim')
            driftservlim=st.session_state.get('driftservlim')

            #st.write("### Información viento")
            normaviento=st.session_state.get('normaviento')
            grupoviento=st.session_state.get('grupoviento')
            tipoviento=st.session_state.get('tipoviento')
            categoriarug=st.session_state.get('categoriarug')
            alturagrad=st.session_state.get('alturagrad')
            exponenteviento=st.session_state.get('exponenteviento')
            facttopoviento=st.session_state.get('facttopoviento')
            tipotopoviento=st.session_state.get('tipotopoviento')

            #st.write("### Software")
            software1=st.session_state.get('software1')
            software2=st.session_state.get('software2')
            software3=st.session_state.get('software3')
            software4=st.session_state.get('software4')
            software5=st.session_state.get('software5')
            software6=st.session_state.get('software6')
            software7=st.session_state.get('software7')

            descripsoft1=st.session_state.get('descripsoft1')
            descripsoft2=st.session_state.get('descripsoft2')
            descripsoft3=st.session_state.get('descripsoft3')
            descripsoft4=st.session_state.get('descripsoft4')
            descripsoft5=st.session_state.get('descripsoft5')
            descripsoft6=st.session_state.get('descripsoft6')
            descripsoft7=st.session_state.get('descripsoft7')

            #st.write("### Modelado")
            tipolosa=st.session_state.get('tipolosa')
            diafragma=st.session_state.get('diafragma')
            crackcolumn=st.session_state.get('crackcolumn')
            crackbeam=st.session_state.get('crackbeam')
            crackwall1=st.session_state.get('crackwall1')
            crackwall2=st.session_state.get('crackwall2')
            crackcoupbeam=st.session_state.get('crackcoupbeam')
            crackslab=st.session_state.get('crackslab')
            rigidpanel=st.session_state.get('rigidpanel')
            amortiguamiento=st.session_state.get('amortiguamiento')

            #st.write("### Cortante basal")
            revcortantebasal=st.session_state.get('revcortantebasal')

            texto_reemplazado = {'proyecto':proyecto, 'fecha':fecha, 'ano':ano, 'mes':mes, 'dia':dia, 'autor':autor, 'inicialesautor':inicialesautor, 'revisor':revisor, 
            'inicialesrevisor':inicialesrevisor, 'direccion':direccion, 'latitud':latitud, 'longitud':longitud, 'descripcionarq':descripcionarq, 'descripciongrav':descripciongrav, 
            'descripsistemlat':descripsistemlat, 'descripsistemcim':descripsistemcim, 'autormecsuelos':autormecsuelos, 'docmecsuelos':docmecsuelos, 'normasismo':normasismo, 
            'regionsismo':regionsismo, 'nivelzonificacion':nivelzonificacion, 'gruposismo':gruposismo, 'descripgruposismo':descripgruposismo, 'programasismo':programasismo, 
            'metodosismo':metodosismo, 'ductilidad':ductilidad, 'driftpclim':driftpclim, 'driftservlim':driftservlim, 'normaviento':normaviento, 'grupoviento':grupoviento, 
            'tipoviento':tipoviento, 'categoriarug':categoriarug, 'alturagrad':alturagrad, 'exponenteviento':exponenteviento, 'facttopoviento':facttopoviento, 
            'tipotopoviento':tipotopoviento, 'software1':software1, 'software2':software2, 'software3':software3, 'software4':software4, 'software5':software5, 'software6':software6, 
            'software7':software7, 'descripsoft1':descripsoft1, 'descripsoft2':descripsoft2, 'descripsoft3':descripsoft3, 'descripsoft4':descripsoft4, 'descripsoft5':descripsoft5,
            'descripsoft6':descripsoft6, 'descripsoft7':descripsoft7, 'tipolosa':tipolosa, 'diafragma':diafragma, 'crackcolumn':crackcolumn, 'crackbeam':crackbeam,
            'crackwall1':crackwall1, 'crackwall2':crackwall2, 'crackcoupbeam':crackcoupbeam, 'crackslab':crackslab, 'rigidpanel':rigidpanel, 'amortiguamiento':amortiguamiento,
            'revcortantebasal':revcortantebasal}

            #st.write(texto_reemplazado)
            
            ####################################REEMPLAZAR VALORES DE TEXTO###########################
            doc=DocxTemplate(template_file)
            doc.render(texto_reemplazado)

            ####################################REEMPLAZAR IMAGENES###########################

            #Introduccion de figura 1.1 Proyecto
            for paragraph in doc.paragraphs:
                if '[Insertar figura1_1]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Insertar figura1_1]', '')
                    paragraph.add_run().add_picture(figura1_1, height=Mm(70))

            #Introduccion de figura 1.2 Ubicacion
            for paragraph in doc.paragraphs:
                if '[Insertar figura1_2]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Insertar figura1_2]', '')
                    paragraph.add_run().add_picture(figura1_2, height=Mm(70))

            #for paragraph in doc.paragraphs:
             #   if '[Insertar figura1_3]' in paragraph.text:
              #      paragraph.text = paragraph.text.replace('[Insertar figura1_3]', '')
               #     paragraph.add_run().add_picture(figuras_arq[0], height=Mm(70))

            #Introduccion de figura 1.3 a figura 1.22 Arquitectura
            for paragraph in doc.paragraphs:
                for i in range(0,num_fig_arq):
                    if f'[Insertar figura1_{i+3}' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'[Insertar figura1_{i+3}]', '')
                        paragraph.add_run().add_picture(figuras_arq[i], height=Mm(70))

            #Introduccion de figura 5.1 plano de cargas
            for paragraph in doc.paragraphs:
                if '[Insertar figura5_1]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Insertar figura5_1]', '')
                    paragraph.add_run().add_picture(figura5_1, height=Mm(70))

             #Introduccion de figura 5.2 regionalizacion
            for paragraph in doc.paragraphs:
                if '[Insertar figura5_2]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Insertar figura5_2]', '')
                    paragraph.add_run().add_picture(figura5_2, height=Mm(70))
            
            #Introduccion de figura 5.3 espectro elastico
            for paragraph in doc.paragraphs:
                if '[Insertar figura5_3]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Insertar figura5_3]', '')
                    paragraph.add_run().add_picture(figura5_3, height=Mm(70))
            
            #Introduccion de figura 5.4 espectro reducido y de servicio
            for paragraph in doc.paragraphs:
                if '[Insertar figura5_4]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Insertar figura5_4]', '')
                    paragraph.add_run().add_picture(figura5_4, height=Mm(70))

            #Introduccion de figura 5.5 a figura 5.9 viento
            for paragraph in doc.paragraphs:
                for i in range(0,num_fig_viento):
                    if f'[Insertar figura5_{i+5}' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'[Insertar figura5_{i+5}]', '')
                        paragraph.add_run().add_picture(figuras_viento[i], height=Mm(70))

            #Introduccion de figura 10.1 a figura 10.5 modelo estructural
            for paragraph in doc.paragraphs:
                for i in range(0,num_fig_modest):
                    if f'[Insertar figura10_{i+1}' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'[Insertar figura10_{i+1}]', '')
                        paragraph.add_run().add_picture(figuras_modest[i], height=Mm(70))

            #Introduccion de figura 10.6 Configuracion modos de vibrar
            for paragraph in doc.paragraphs:
                if '[Insertar figura10_6a]' and '[Insertar figura10_6b]' and '[Insertar figura10_6c]' in paragraph.text:
                    paragraph.text = paragraph.text.replace('[Insertar figura10_6a]', '')
                    paragraph.text = paragraph.text.replace('[Insertar figura10_6b]', '')
                    paragraph.text = paragraph.text.replace('[Insertar figura10_6c]', '')
                    paragraph.add_run().add_picture(figura10_6a, width=Mm(50))
                    paragraph.add_run().add_picture(figura10_6b, width=Mm(50))
                    paragraph.add_run().add_picture(figura10_6c, width=Mm(50))
            
            #Introduccion de figura 10.9 a figura 10.28 deflexiones verticales
            for paragraph in doc.paragraphs:
                for i in range(0,num_fig_dz):
                    if f'[Insertar figura10_{i+9}' in paragraph.text:
                        paragraph.text = paragraph.text.replace(f'[Insertar figura10_{i+9}]', '')
                        paragraph.add_run().add_picture(figuras_dz[i], height=Mm(70))


            def generar_word():
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer
            buffer = generar_word()

            if st.button("Paso 9"):
                go_to_step(9)
        
            elif st.download_button(
                    label="Descargar memoria",
                    data=buffer,
                    file_name="Memoria_calculo.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                st.success("¡Proceso completado!") 
             
                
    else: 
        st.warning("No se ha cargado plantilla")

main()
